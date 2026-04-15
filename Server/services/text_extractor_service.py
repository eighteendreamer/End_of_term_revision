"""
文本提取服务
从不同格式的文件中提取文本内容
"""
import io
import logging
from typing import Optional
import pdfplumber
from docx import Document
from PIL import Image

logger = logging.getLogger(__name__)


class TextExtractorService:
    """文本提取服务类"""
    
    @staticmethod
    def extract_from_pdf(file_data: bytes) -> str:
        """
        从PDF文件提取文本
        :param file_data: PDF文件字节数据
        :return: 提取的文本
        """
        try:
            text_content = []
            
            with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            extracted_text = '\n\n'.join(text_content)
            
            if not extracted_text.strip():
                raise ValueError("PDF文件中未提取到文本内容")
            
            logger.info(f"成功从PDF提取文本，长度: {len(extracted_text)}")
            return extracted_text
            
        except Exception as e:
            logger.error(f"PDF文本提取失败: {str(e)}")
            raise ValueError(f"PDF文本提取失败: {str(e)}")
    
    @staticmethod
    def extract_from_word(file_data: bytes) -> str:
        """
        从Word文档提取文本
        :param file_data: Word文件字节数据
        :return: 提取的文本
        """
        try:
            doc = Document(io.BytesIO(file_data))
            
            text_content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)
            
            extracted_text = '\n\n'.join(text_content)
            
            if not extracted_text.strip():
                raise ValueError("Word文档中未提取到文本内容")
            
            logger.info(f"成功从Word提取文本，长度: {len(extracted_text)}")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Word文本提取失败: {str(e)}")
            raise ValueError(f"Word文本提取失败: {str(e)}")
    
    @staticmethod
    def extract_from_text(file_data: bytes) -> str:
        """
        从文本文件提取内容
        :param file_data: 文本文件字节数据
        :return: 文本内容
        """
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
            
            for encoding in encodings:
                try:
                    text = file_data.decode(encoding)
                    if text.strip():
                        logger.info(f"成功使用{encoding}编码读取文本，长度: {len(text)}")
                        return text
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("无法识别文本文件编码")
            
        except Exception as e:
            logger.error(f"文本文件读取失败: {str(e)}")
            raise ValueError(f"文本文件读取失败: {str(e)}")
    
    @staticmethod
    def extract_from_image(file_data: bytes, use_ai: bool = False) -> str:
        """
        从图片提取文本（OCR）
        :param file_data: 图片文件字节数据
        :param use_ai: 是否使用AI进行OCR（需要配置AI服务）
        :return: 提取的文本
        """
        try:
            # 验证图片可以打开
            image = Image.open(io.BytesIO(file_data))
            width, height = image.size
            
            logger.info(f"图片尺寸: {width}x{height}")
            
            if use_ai:
                # TODO: 集成AI视觉模型进行OCR
                # 可以使用OpenAI GPT-4 Vision或其他OCR服务
                raise NotImplementedError("AI OCR功能尚未实现")
            else:
                # 暂时返回提示信息
                return f"[图片文件: {width}x{height}像素]\n\n注意: 图片OCR功能需要配置AI服务。当前仅保存图片信息。"
            
        except NotImplementedError:
            raise
        except Exception as e:
            logger.error(f"图片处理失败: {str(e)}")
            raise ValueError(f"图片处理失败: {str(e)}")
    
    @staticmethod
    def extract_text(file_data: bytes, file_type: str, use_ai_ocr: bool = False) -> str:
        """
        根据文件类型提取文本
        :param file_data: 文件字节数据
        :param file_type: 文件类型（pdf/docx/txt/jpg等）
        :param use_ai_ocr: 是否使用AI进行图片OCR
        :return: 提取的文本
        """
        file_type = file_type.lower()
        
        try:
            if file_type == 'pdf':
                return TextExtractorService.extract_from_pdf(file_data)
            
            elif file_type in ['doc', 'docx']:
                return TextExtractorService.extract_from_word(file_data)
            
            elif file_type in ['txt', 'md']:
                return TextExtractorService.extract_from_text(file_data)
            
            elif file_type in ['jpg', 'jpeg', 'png']:
                return TextExtractorService.extract_from_image(file_data, use_ai_ocr)
            
            else:
                raise ValueError(f"不支持的文件类型: {file_type}")
                
        except Exception as e:
            logger.error(f"文本提取失败 ({file_type}): {str(e)}")
            raise
